import streamlit as st
import os
import time
import re
from io import BytesIO
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from docx import Document
from openai import AzureOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LDocument
from Modules.prompts import (
    get_executive_summary_and_objective_prompt,
    get_scope_prereq_assumptions_prompt,     get_resource_schedule_and_commercial_prompt, get_communication_plan_prompt
)


# -------------------------------------------------------
# 1. SETUP
# -------------------------------------------------------
load_dotenv()
KNOWLEDGE_FOLDER = "Knowledge_Repo"
PERSIST_DIR = "chroma_db"

st.set_page_config(page_title="📄 RFP Executive Summary Generator", layout="wide")
st.markdown("<h1 style='text-align:center; color:#4B0082;'>📄 Executive Summary Generator</h1>", unsafe_allow_html=True)
st.markdown("---")

# -------------------------------------------------------
# 2. UTILITIES
# -------------------------------------------------------

def extract_text(file):
    """Extract text from PDF or DOCX"""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def build_knowledge_base(folder=KNOWLEDGE_FOLDER, persist_dir=PERSIST_DIR):
    """Build or load persistent Chroma vector DB from Knowledge_Repo"""
    os.makedirs(folder, exist_ok=True)
    os.makedirs(persist_dir, exist_ok=True)

    embedding_model = AzureOpenAIEmbeddings(
        model="text-embedding-ada-002",
        azure_endpoint=os.getenv("AZURE_OPENAI_EMD_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_EMD_KEY"),
        api_version=os.getenv("AZURE_OPENAI_EMD_VERSION")
    )

    if os.listdir(persist_dir):
        return Chroma(
            embedding_function=embedding_model,
            persist_directory=persist_dir,
            collection_name="rfp_responses"
        )

    docs = []
    for f in os.listdir(folder):
        if f.endswith((".pdf", ".docx")):
            path = os.path.join(folder, f)
            text = extract_text(open(path, "rb"))
            if text.strip():
                docs.append(LDocument(page_content=text, metadata={"source": f}))

    if not docs:
        raise ValueError(f"No readable files found in {folder}")

    return Chroma.from_documents(
        documents=docs,
        embedding=embedding_model,
        persist_directory=persist_dir,
        collection_name="rfp_responses"
    )

from docx import Document


def insert_executive_summary_into_template(
    template_path,
    summary_text,
    objective_text=None,
    scope_text=None,
    resource_schedule_text=None,
    communication_plan_text=None
):
    """
    Replace placeholders in the template:
    <<EXEC_SUMMARY>>, <<OBJECTIVE>>, <<SCOPE_TEXT>>, <<RESOURCE_SCHEDULE>>
    """

    doc = Document(template_path)

    def replace_placeholder(doc, placeholder, new_text):
        if not new_text:
            return
        for para in doc.paragraphs:
            if placeholder in "".join(run.text for run in para.runs):
                para.clear()
                para.text = ""
                parent = para._element.getparent()
                idx = parent.index(para._element)

                for line in new_text.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        new_para = doc.add_paragraph()
                    elif stripped.startswith("###"):
                        bold_text = stripped.lstrip("#").strip()
                        new_para = doc.add_paragraph()
                        run = new_para.add_run(bold_text)
                        run.bold = True
                    elif stripped.startswith("•"):
                        new_para = doc.add_paragraph(stripped.lstrip("•").strip(), style="List Bullet")
                    else:
                        new_para = doc.add_paragraph(stripped)
                    parent.insert(idx + 1, new_para._element)
                    idx += 1

                parent.remove(para._element)
                return

    # Replace placeholders
    replace_placeholder(doc, "<<EXEC_SUMMARY>>", summary_text)
    replace_placeholder(doc, "<<OBJECTIVE>>", objective_text)
    replace_placeholder(doc, "<<SCOPE_TEXT>>", scope_text)
    replace_placeholder(doc, "<<RESOURCE_SCHEDULE>>", resource_schedule_text)
    replace_placeholder(doc, "<<COMMUNICATION_PLAN>>", communication_plan_text)


    return doc



def generate_exec_summary_and_objective(reference_text, condensed_rfp, num_interfaces=113):
    """Generate both sections and return as separate strings."""
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_executive_summary_and_objective_prompt(reference_text, condensed_rfp, num_interfaces)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    full_output = response.choices[0].message.content.strip()

    # --- Split into Executive Summary and Objective ---
    exec_match = re.search(r"\*\*?Executive Summary\*\*?\s*(.*?)\s*(?=\*\*?Objective\*\*?)", full_output, re.S | re.I)
    obj_match = re.search(r"\*\*?Objective\*\*?\s*(.*)", full_output, re.S | re.I)

    exec_text = exec_match.group(1).strip() if exec_match else full_output
    obj_text = obj_match.group(1).strip() if obj_match else ""

    return exec_text, obj_text

def generate_scope_sections(reference_text, condensed_rfp, num_interfaces=None):
    """Generate combined 'In Scope', 'Prerequisites', 'Assumptions', and 'Out of Scope' section as one block."""
    # from openai import AzureOpenAI
    # import os
    # from Modules.prompts import get_scope_prereq_assumptions_prompt

    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_scope_prereq_assumptions_prompt(reference_text, condensed_rfp, num_interfaces)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

def generate_resource_schedule_and_commercial(reference_text, condensed_rfp):
    """Generate Resource Schedule and Commercials section."""
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_resource_schedule_and_commercial_prompt(reference_text, condensed_rfp)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()
def generate_communication_plan(reference_text, condensed_rfp):
    """Generate Communication Plan section."""
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    prompt = get_communication_plan_prompt(reference_text, condensed_rfp)
    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()


# -------------------------------------------------------
# 3. STREAMLIT UI
# -------------------------------------------------------

st.subheader("Step 1: Upload RFP Document")
uploaded_file = st.file_uploader("Choose PDF or DOCX", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("🔍 Extracting RFP content..."):
        rfp_text = extract_text(uploaded_file)
        time.sleep(2)
    st.success("✅ RFP content extracted!")

    st.text_area("Preview RFP Text", rfp_text[:2000] + "..." if len(rfp_text) > 2000 else rfp_text, height=250)

    st.subheader("Step 2: Build or Load Knowledge Base")
    with st.spinner("📚 Loading Knowledge Repository..."):
        knowledge_db = build_knowledge_base()
        time.sleep(2)
    st.success("✅ Knowledge base ready (persistent)!")

    with st.spinner("🔎 Retrieving relevant reference material..."):
        retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
        # ref_docs = retriever.get_relevant_documents(rfp_text)
        ref_docs = retriever.invoke(rfp_text)
        reference_text = "\n\n".join([d.page_content for d in ref_docs])
    st.success(f"✅ Retrieved {len(ref_docs)} relevant reference documents!")

    st.subheader("Step 3: Generate Core Proposal Section")
    st.markdown("This will generate both the Executive Summary and Objective sections using Crave InfoTech’s reference knowledge.")

    num_interfaces = st.number_input(
        "Enter number of interfaces (if known):",
        min_value=0, max_value=10000, value=113, step=1,
        help="Used to reference the migration scale dynamically."
    )

    with st.spinner("✍️ Generating proposal content..."):
        exec_summary, objective = generate_exec_summary_and_objective(reference_text, rfp_text,num_interfaces)

    st.success("✅ Executive Summary and Objective generated successfully!")


    # ------------------------------
    # Step 4: Generate Scope Sections
    # ------------------------------
    st.subheader("Step 4: Generate Scope, Prerequisites, Assumptions, and Out of Scope")

    with st.spinner("🧩 Generating detailed scope and assumptions..."):
        scope_text = generate_scope_sections(reference_text, rfp_text, num_interfaces)
    st.success("✅ Scope and Assumptions section generated!")

    # ------------------------------
    # Step 5: Generate Resource Schedule & Commercials
    # ------------------------------
    st.subheader("Step 5: Generate Resource Schedule & Commercials")

    with st.spinner("📊 Generating Resource Schedule and Commercials..."):
        resource_schedule_text = generate_resource_schedule_and_commercial(reference_text, rfp_text)
    st.success("✅ Resource Schedule and Commercials generated successfully!")

    # ------------------------------
    # Step 6: Generate communication_plan_text
    # ------------------------------

    with st.spinner("📢 Generating Communication Plan..."):
        communication_plan_text = generate_communication_plan(reference_text, rfp_text)
    st.success("✅ Communication Plan generated successfully!")

    # ------------------------------
    # Step 4 + 5: Generate Scope + Resource Schedule in Parallel
    # ------------------------------
    # st.subheader("Step 4: Generate Scope and Resource Schedule")

    # with st.spinner("🧩 Generating Scope, Assumptions, and Resource Schedule in parallel..."):
    #     with concurrent.futures.ThreadPoolExecutor() as executor:
    #         f1 = executor.submit(generate_scope_sections, reference_text, rfp_text, num_interfaces)
    #         f2 = executor.submit(generate_resource_schedule_and_commercial, reference_text, rfp_text)
    #         scope_text = f1.result()
    #         resource_schedule_text = f2.result()

    # st.success("✅ Scope, Assumptions, and Resource Schedule generated successfully!")

    
    # -------------------------------------------------------
    # Step 6: Insert into Template & Download
    # -------------------------------------------------------
    st.subheader("Step 7: Insert into Template & Download")
    template_path = "Template/PIPO TO IS Response Template.docx"

    if not os.path.exists(template_path):
        st.error(f"Template not found at {template_path}")
    else:
        final_doc = insert_executive_summary_into_template(
            template_path,
            summary_text=exec_summary,
            objective_text=objective,
            scope_text=scope_text,
            resource_schedule_text=resource_schedule_text,
            communication_plan_text=communication_plan_text
        )


        buffer = BytesIO()
        final_doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Final Proposal (DOCX)",
            data=buffer,
            file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

